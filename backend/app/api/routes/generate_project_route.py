# backend/app/api/routes/generate_extended_project_route.py

from fastapi import APIRouter, HTTPException, BackgroundTasks, Depends
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from typing import Dict, Any, Optional, List, TypedDict, Annotated, Union
import logging
import uuid
import os
from dotenv import load_dotenv
import asyncio
import json
import re # For potential reference parsing if needed

# Langchain & LangGraph specific imports
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langchain_core.messages import HumanMessage, AIMessage, BaseMessage, ToolMessage, SystemMessage
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_core.runnables import RunnableConfig, RunnablePassthrough
from langchain_anthropic import ChatAnthropic

# --- Environment Setup ---
load_dotenv()
# Check for API keys
anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
tavily_api_key = os.getenv("TAVILY_API_KEY")

# --- Logging Setup ---
logger = logging.getLogger(__name__)
# Configure logging level and format (adjust as needed)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - [%(levelname)s] - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)


# --- API Router Setup ---
router = APIRouter()

# --- Pydantic Models for API Request/Response ---
class ExtendedProjectInput(BaseModel):
    field: str = Field(..., description="The general field of the project (e.g., 'Artificial Intelligence', 'Renewable Energy').")
    domain: str = Field(..., description="The specific domain within the field (e.g., 'Natural Language Processing', 'Solar Panel Efficiency').")
    idea: Optional[str] = Field(None, description="An initial idea or specific focus for the project (optional).")
    academic_level: str = Field("undergraduate", description="Academic level: 'undergraduate', 'masters', or 'phd'")
    page_target: int = Field(7, ge=5, le=15, description="Target number of pages (5-15 recommended)") # Added range validation
    format_type: str = Field("ieee", description="Format type: 'ieee', 'acm', 'apa', etc.")

class ExtendedProposalResponse(BaseModel):
    proposal_id: str
    status: str
    message: str
    result: Optional[Dict[str, Any]] = None

# --- LangGraph State Definition ---
class ExtendedProposalGraphState(TypedDict):
    """
    Represents the state of the extended project proposal generation graph.
    """
    # Inputs
    field: str
    domain: str
    idea: Optional[str]
    academic_level: str
    page_target: int
    format_type: str

    # Node outputs / Intermediate results (more granular)
    problem_statement: Optional[str]
    problem_background: Optional[str]
    project_significance: Optional[str]
    key_challenges: Optional[List[str]]
    search_queries: Optional[List[str]]
    search_results: Optional[List[Dict[str, str]]] # Store raw search hits
    literature_summary: Optional[str]
    literature_detailed_review: Optional[Dict[str, str]]  # Section title -> content map
    research_gaps: Optional[List[str]]
    research_questions: Optional[List[str]]
    project_objectives: Optional[List[str]]

    # Methodology sections
    methodology_approach_overview: Optional[str] # Renamed for clarity
    methodology_sections: Optional[Dict[str, str]]  # Section title -> content map (e.g., Data Collection, Analysis)
    experimental_design: Optional[str] # Switched to string for LLM generation flexibility
    methodology_tools: Optional[List[str]]
    validation_strategy: Optional[str] # Renamed for clarity

    # Results and Discussion
    expected_results_summary: Optional[str] # Summary first
    expected_results_detailed: Optional[Dict[str, str]]  # Section title -> content map
    discussion_sections: Optional[Dict[str, str]] # Section title -> content map (e.g., Implications, Comparison)
    expected_metrics: Optional[List[str]]
    expected_impact: Optional[str]
    limitations_and_future_work: Optional[str]

    # Supporting content
    timeline_and_milestones: Optional[str] # String from LLM
    # budget_and_resources: Optional[str] # Added optional budget section

    # References and formatting
    references_data: Optional[List[Dict[str, str]]] # Store structured ref info if possible
    formatted_references: Optional[str] # Final formatted list
    executive_summary: Optional[str]  # Abstract/summary

    # Final Output
    formatted_proposal: Optional[str]  # Full proposal document

    # Control Flow & Error Handling
    error: Optional[str] = None
    node_history: List[str] = Field(default_factory=list)

# --- Tool Definition ---
search_tool = None
if tavily_api_key:
    search_tool = TavilySearchResults(max_results=7) # Increase results slightly

# --- LLM Definition ---
# Using Claude 3.5 Sonnet model
llm = None
if anthropic_api_key:
    llm = ChatAnthropic(model="claude-3-5-sonnet-20240620", temperature=0.9) # Slightly higher temp for creativity

# --- Helper Function for JSON Parsing ---
async def parse_llm_json_output(llm_call_coroutine, expected_keys: List[str], node_name: str) -> Union[Dict, str]:
    """Attempts to invoke LLM, parse JSON, and validate keys. Returns dict or error string."""
    if llm is None:
        logger.error(f"[{node_name}] LLM not initialized - API key might be missing")
        return f"LLM not initialized in {node_name}. API key might be missing."
        
    try:
        response_raw = await llm_call_coroutine
        # Attempt to extract JSON even if markdown code block ```json ... ``` is present
        json_match = re.search(r"```json\s*([\s\S]*?)\s*```", response_raw.content, re.IGNORECASE)
        if json_match:
            json_str = json_match.group(1)
        else:
            # Assume raw content is JSON or attempt parsing anyway
            json_str = response_raw.content

        parsed_json = json.loads(json_str)

        if not isinstance(parsed_json, dict):
            raise ValueError("LLM output was not a valid JSON object.")

        missing_keys = [key for key in expected_keys if key not in parsed_json]
        if missing_keys:
            logger.warning(f"[{node_name}] LLM JSON response missing keys: {missing_keys}. Content: {json_str[:500]}")
            # Decide whether to error out or proceed with partial data
            # For now, proceed but log warning. Could add error handling here.

        return parsed_json
    except json.JSONDecodeError as e:
        logger.error(f"[{node_name}] Failed to decode LLM JSON response: {e}. Raw content: {response_raw.content[:500]}")
        return f"Failed to decode JSON response in {node_name}"
    except Exception as e:
        logger.error(f"[{node_name}] Error processing LLM response: {e}. Raw content: {getattr(response_raw, 'content', str(response_raw))[:500]}")
        return f"Error processing LLM response in {node_name}: {str(e)}"

# --- Node Functions (Enhanced) ---

async def initialize_project_context(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 1: Creates detailed problem context, significance, challenges, RQs, objectives.
    """
    node_name = "initialize_project_context"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    prompt_template = """You are a senior professor advising a {academic_level} student in {field} ({domain}).
Develop the initial sections for a comprehensive Final Year Project proposal.

Based on:
Field: {field}
Domain: {domain}
Optional Idea: {idea}

Generate the following, aiming for depth suitable for a {page_target}-page proposal:

1.  **Problem Statement (approx. 300-400 words):** Detailed articulation of the problem, context, and relevance.
2.  **Problem Background (approx. 400-500 words):** Elaborate on the history, prior work, and factors leading to the problem.
3.  **Project Significance (approx. 250-350 words):** Explain academic/practical importance and potential benefits.
4.  **Key Challenges (approx. 300-400 words):** Describe 4-6 major hurdles (technical, theoretical, practical).
5.  **Research Questions (3-5 specific questions):** Formulate precise questions the project aims to answer.
6.  **Project Objectives (3-5 measurable objectives):** Define clear goals for successful project completion.

Return a JSON object with keys: "problem_statement", "problem_background", "project_significance", "key_challenges", "research_questions", "project_objectives".
"""
    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | llm # Expecting JSON output directly based on instructions

    invoke_coro = chain.ainvoke({
        "academic_level": state["academic_level"],
        "field": state["field"],
        "domain": state["domain"],
        "idea": state.get("idea") or "General advancements in this area",
        "page_target": state["page_target"]
    }, config=config)

    expected_keys = ["problem_statement", "problem_background", "project_significance", "key_challenges", "research_questions", "project_objectives"]
    result = await parse_llm_json_output(invoke_coro, expected_keys, node_name)

    if isinstance(result, str): # Error occurred
        return {"error": result}
    else:
        logger.info(f"[{node_name}] Initial context generated successfully.")
        # Update state with potentially missing keys set to default values
        return {
            "problem_statement": result.get("problem_statement"),
            "problem_background": result.get("problem_background"),
            "project_significance": result.get("project_significance"),
            "key_challenges": result.get("key_challenges", []),
            "research_questions": result.get("research_questions", []),
            "project_objectives": result.get("project_objectives", []),
            "error": None
         }

async def comprehensive_literature_review(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 2: Conducts extensive literature search and synthesizes a detailed review with sections and gaps.
    """
    node_name = "comprehensive_literature_review"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    problem = state.get("problem_statement")
    if not problem: return {"error": "Problem statement missing for literature review."}

    # 1. Generate Search Queries
    query_prompt = ChatPromptTemplate.from_template(
        "Generate 6-8 specific search queries for a thorough literature review on: Problem: '{problem}'. Focus on foundational work, recent advances, methodologies, challenges, and related fields. Output as JSON list."
    )
    query_chain = query_prompt | llm | JsonOutputParser()
    try:
        search_queries = await query_chain.ainvoke({"problem": problem[:500]}, config=config) # Limit problem length
        if not isinstance(search_queries, list) or len(search_queries) < 4:
            raise ValueError("Invalid query list generated")
        logger.info(f"[{node_name}] Generated {len(search_queries)} search queries.")
    except Exception as e:
        logger.error(f"[{node_name}] Error generating search queries: {e}, using defaults.")
        search_queries = [ f"{state['field']} {state['domain']} review", f"{problem[:100]} solutions research", f"{state['domain']} challenges methods" ]

    # 2. Execute Search (Async for potentially multiple queries)
    search_tasks = [search_tool.ainvoke(query, config=config) for query in search_queries[:6]] # Limit concurrent searches
    search_results_raw = await asyncio.gather(*search_tasks, return_exceptions=True)

    all_results = []
    references_data = []
    processed_urls = set() # Avoid duplicates
    for result in search_results_raw:
        if isinstance(result, Exception):
            logger.warning(f"[{node_name}] Search query failed: {result}")
            continue
        if isinstance(result, list):
            for item in result:
                 if isinstance(item, dict) and item.get('url') and item.get('url') not in processed_urls:
                    content_snippet = item.get('content', '')[:500] # Limit snippet length stored
                    ref_info = {
                        "title": item.get('title', 'Untitled Source'),
                        "url": item.get('url'),
                        "content": content_snippet, # Store snippet for context
                        # Placeholder fields, ideally populate these if search tool provides them or via another step
                        "authors": item.get('authors', 'Unknown Authors'),
                        "year": item.get('year', 'N/A'),
                        "journal": item.get('source', 'Unknown Source')
                    }
                    all_results.append(ref_info) # Use the structured dict
                    references_data.append(ref_info)
                    processed_urls.add(item['url'])


    logger.info(f"[{node_name}] Processed {len(all_results)} unique search results.")
    if not all_results:
        logger.warning(f"[{node_name}] No search results found. Literature review will rely heavily on LLM's knowledge.")
        # Decide if this should be an error or just proceed. Proceeding for now.

    # 3. Synthesize Detailed Literature Review via LLM
    review_prompt_template = """You are an expert academic writer synthesizing a literature review for a {academic_level} project in {field} ({domain}).

Problem: {problem}
Background: {background}
Challenges: {challenges}
Identified Research Questions: {research_questions}
Search Result Snippets (use for context, supplement with domain knowledge):
{search_snippets}

**Task:** Generate a comprehensive literature review (target 1000-1500 words total) structured as follows:

1.  **Literature Summary (approx. 300 words):** Overview of the current state-of-the-art.
2.  **Detailed Review Sections (4-6 sections, approx. 200-300 words each):** Create distinct sections covering different facets (e.g., 'Foundational Theories', 'Current Methodologies', 'Key Empirical Studies', 'Identified Limitations in Prior Work'). Analyze relevant work within each section.
3.  **Identified Research Gaps (5-7 specific gaps):** Clearly articulate gaps based on the review, linking them to the project's problem and RQs.

Return ONLY a JSON object with keys: "literature_summary" (string), "literature_detailed_review" (object: section_title -> content), "research_gaps" (list of strings).
"""
    search_snippets_str = "\n---\n".join([f"Title: {r['title']}\nURL: {r['url']}\nSnippet: {r['content']}" for r in all_results[:10]]) # Provide context from top hits
    prompt = ChatPromptTemplate.from_template(review_prompt_template)
    chain = prompt | llm # Expect JSON

    invoke_coro = chain.ainvoke({
        "academic_level": state["academic_level"],
        "field": state["field"],
        "domain": state["domain"],
        "problem": state.get("problem_statement", ""),
        "background": state.get("problem_background", ""),
        "challenges": state.get("key_challenges", []),
        "research_questions": state.get("research_questions", []),
        "search_snippets": search_snippets_str or "No specific search results found. Synthesize based on general knowledge."
    }, config=config)

    expected_keys = ["literature_summary", "literature_detailed_review", "research_gaps"]
    result = await parse_llm_json_output(invoke_coro, expected_keys, node_name)

    if isinstance(result, str):
        return {"error": result}
    else:
        logger.info(f"[{node_name}] Detailed literature review generated.")
        # Store raw search results and structured references separately
        return {
            "search_queries": search_queries,
            "search_results": all_results, # Store processed results
            "references_data": references_data, # Store data for reference formatting
            "literature_summary": result.get("literature_summary"),
            "literature_detailed_review": result.get("literature_detailed_review", {}),
            "research_gaps": result.get("research_gaps", []),
            "error": None
        }


async def detailed_methodology_design(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 3: Designs a detailed methodology chapter with multiple sections.
    """
    node_name = "detailed_methodology_design"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    problem = state.get("problem_statement")
    objectives = state.get("project_objectives")
    gaps = state.get("research_gaps")
    questions = state.get("research_questions")
    if not problem or not objectives or not gaps or not questions:
        return {"error": "Missing problem, objectives, gaps, or questions for methodology design."}

    methodology_prompt_template = """You are a Research Methodologist designing the methodology for a {academic_level} project in {field} ({domain}).

Problem: {problem}
Research Questions: {questions}
Objectives: {objectives}
Gaps Addressed: {gaps}

**Task:** Create a detailed Methodology chapter (target 1200-1800 words total) with the following structure:

1.  **Methodology Overview (approx. 300 words):** Describe and justify the chosen research approach (e.g., quantitative, qualitative, mixed-methods, design science).
2.  **Research Design (approx. 300 words):** Detail the specific design (e.g., experimental, quasi-experimental, case study, survey).
3.  **Data Collection (approx. 300 words):** Specify data sources, instruments, and procedures. If developing software/system, describe development process.
4.  **Data Analysis Plan (approx. 300 words):** Outline statistical tests, algorithms, coding procedures, or analytical techniques.
5.  **Validation Strategy (approx. 200 words):** Explain how results/artifacts will be validated (metrics, comparisons, user testing).
6.  **Tools & Technologies (list):** Specify software, libraries, hardware, datasets required.
7.  **Ethical Considerations (approx. 150 words, if applicable):** Discuss relevant ethical issues and mitigation.

Return ONLY a JSON object with keys: "methodology_overview", "research_design", "data_collection", "data_analysis_plan", "validation_strategy", "tools_and_technologies", "ethical_considerations". The value for each key (except tools) should be the generated text string. 'tools_and_technologies' should be a list of strings.
"""
    prompt = ChatPromptTemplate.from_template(methodology_prompt_template)
    chain = prompt | llm # Expect JSON

    invoke_coro = chain.ainvoke({
        "academic_level": state["academic_level"],
        "field": state["field"],
        "domain": state["domain"],
        "problem": problem,
        "questions": questions,
        "objectives": objectives,
        "gaps": gaps
    }, config=config)

    expected_keys = ["methodology_overview", "research_design", "data_collection", "data_analysis_plan", "validation_strategy", "tools_and_technologies", "ethical_considerations"]
    result = await parse_llm_json_output(invoke_coro, expected_keys, node_name)

    if isinstance(result, str):
        return {"error": result}
    else:
        logger.info(f"[{node_name}] Detailed methodology generated.")
        # Structure the output for the state
        methodology_sections = {
            "Research Design": result.get("research_design"),
            "Data Collection": result.get("data_collection"),
            "Data Analysis Plan": result.get("data_analysis_plan"),
            "Ethical Considerations": result.get("ethical_considerations")
            # Add more if the prompt generated them differently
        }
        # Filter out None values
        methodology_sections = {k: v for k, v in methodology_sections.items() if v}

        return {
            "methodology_approach_overview": result.get("methodology_overview"),
            "methodology_sections": methodology_sections,
            "validation_strategy": result.get("validation_strategy"),
            "methodology_tools": result.get("tools_and_technologies", []),
            "error": None
        }


async def generate_results_discussion_and_support(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 4: Generates detailed Expected Results, Discussion, Impact, Limitations, Future Work, and Timeline.
    """
    node_name = "generate_results_discussion_and_support"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    problem = state.get("problem_statement")
    objectives = state.get("project_objectives")
    questions = state.get("research_questions")
    methodology_overview = state.get("methodology_approach_overview")
    validation_strategy = state.get("validation_strategy")

    if not problem or not objectives or not questions or not methodology_overview or not validation_strategy:
        return {"error": "Missing prerequisites for Results & Discussion generation."}

    results_prompt_template = """You are projecting outcomes and discussion points for a {academic_level} research project in {field} ({domain}).

Problem: {problem}
Research Questions: {questions}
Objectives: {objectives}
Methodology Overview: {methodology_overview}
Validation Strategy: {validation_strategy}

**Task:** Generate content for the Expected Results, Discussion, and Supporting sections (target 1500-2000 words total):

1.  **Expected Results Summary (approx. 300 words):** High-level overview of anticipated findings related to objectives.
2.  **Detailed Expected Results (3-5 sections, approx. 200-250 words each):** Project specific outcomes for key objectives or research questions. Describe potential data patterns, metrics, or artifact characteristics.
3.  **Discussion (3-5 sections, approx. 200-250 words each):** Interpret the expected results, discuss implications, compare with literature (hypothetically), address significance. Sections could cover 'Interpretation', 'Implications', 'Comparison to Prior Work'.
4.  **Expected Impact (approx. 300 words):** Detail scientific, practical, or societal contributions.
5.  **Limitations and Future Work (approx. 300 words):** Acknowledge potential limitations and suggest subsequent research directions.
6.  **Project Timeline (detailed phases/milestones):** Create a realistic timeline with key milestones (e.g., Literature Review: Month 1-2, Development: Month 3-5, etc.).
7.  **Expected Metrics (list 5-7):** List specific metrics from the validation strategy.

Return ONLY a JSON object with keys: "expected_results_summary", "detailed_expected_results" (object: section_title -> content), "discussion" (object: section_title -> content), "expected_impact", "limitations_and_future_work", "project_timeline" (string description), "expected_metrics" (list of strings).
"""
    prompt = ChatPromptTemplate.from_template(results_prompt_template)
    chain = prompt | llm # Expect JSON

    invoke_coro = chain.ainvoke({
        "academic_level": state["academic_level"],
        "field": state["field"],
        "domain": state["domain"],
        "problem": problem,
        "questions": questions,
        "objectives": objectives,
        "methodology_overview": methodology_overview,
        "validation_strategy": validation_strategy
    }, config=config)

    expected_keys = ["expected_results_summary", "detailed_expected_results", "discussion", "expected_impact", "limitations_and_future_work", "project_timeline", "expected_metrics"]
    result = await parse_llm_json_output(invoke_coro, expected_keys, node_name)

    if isinstance(result, str):
        return {"error": result}
    else:
        logger.info(f"[{node_name}] Results, Discussion, and Support sections generated.")
        return {
            "expected_results_summary": result.get("expected_results_summary"),
            "expected_results_detailed": result.get("detailed_expected_results", {}),
            "discussion_sections": result.get("discussion", {}), # Map 'discussion' key from LLM
            "expected_impact": result.get("expected_impact"),
            "limitations_and_future_work": result.get("limitations_and_future_work"),
            "timeline_and_milestones": result.get("project_timeline"), # Store as string
            "expected_metrics": result.get("expected_metrics", []),
            "error": None
        }

async def generate_executive_summary(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 5: Generates the Executive Summary/Abstract.
    """
    node_name = "generate_executive_summary"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    # Gather key info generated so far
    problem = state.get("problem_statement", "N/A")
    objectives = state.get("project_objectives", [])
    methodology = state.get("methodology_approach_overview", "N/A")
    results_summary = state.get("expected_results_summary", "N/A")
    impact = state.get("expected_impact", "N/A")

    if "N/A" in [problem, methodology, results_summary, impact]:
        return {"error": "Cannot generate summary, missing core proposal components."}

    summary_prompt_template = """Generate a concise Executive Summary / Abstract (approx. 300-400 words) for this project proposal. It should encapsulate:
- The core problem and motivation.
- The main objectives.
- The proposed methodology approach.
- The key expected outcomes and their potential impact.

Maintain a formal, academic tone.

Problem: {problem}
Objectives: {objectives}
Methodology: {methodology}
Expected Results Summary: {results_summary}
Expected Impact: {impact}

Output ONLY the summary text.
"""
    prompt = ChatPromptTemplate.from_template(summary_prompt_template)
    chain = prompt | llm | StrOutputParser()

    try:
        summary = await chain.ainvoke({
            "problem": problem,
            "objectives": "\n- ".join(objectives),
            "methodology": methodology,
            "results_summary": results_summary,
            "impact": impact
        }, config=config)
        logger.info(f"[{node_name}] Executive Summary generated.")
        return {"executive_summary": summary, "error": None}
    except Exception as e:
        logger.error(f"[{node_name}] Error generating summary: {e}")
        return {"error": f"Failed to generate Executive Summary: {str(e)}"}

async def format_references_node(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 6: Formats the collected references according to the specified style.
    """
    node_name = "format_references_node"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    references_data = state.get("references_data", [])
    format_type = state.get("format_type", "ieee").upper() # Ensure uppercase for prompt

    if not references_data:
        logger.warning(f"[{node_name}] No reference data found from search. Will attempt to generate placeholders.")
        # Optional: Add LLM call here to generate plausible references based on field/domain if desired
        return {"formatted_references": "References section could not be generated due to lack of source data.", "error": None}


    # Prepare data for LLM formatting
    refs_input_list = []
    for i, ref in enumerate(references_data[:15]): # Limit refs sent to LLM
        refs_input_list.append(f"Reference {i+1}:\nTitle: {ref.get('title', 'N/A')}\nURL: {ref.get('url', 'N/A')}\nAuthors: {ref.get('authors', 'N/A')}\nYear: {ref.get('year', 'N/A')}\nSource: {ref.get('journal', 'N/A')}")

    refs_input_str = "\n\n".join(refs_input_list)

    formatting_prompt_template = """You are an expert citation formatter. Format the following references according to the {format_type} citation style.
Ensure correct numbering, punctuation, and ordering as per the style guide.

Raw Reference Data:
{raw_references}

Output ONLY the formatted reference list, numbered sequentially, with each reference on a new line.
"""
    prompt = ChatPromptTemplate.from_template(formatting_prompt_template)
    chain = prompt | llm | StrOutputParser()

    try:
        formatted_refs = await chain.ainvoke({
            "format_type": format_type,
            "raw_references": refs_input_str
        }, config=config)
        logger.info(f"[{node_name}] References formatted successfully in {format_type} style.")
        return {"formatted_references": formatted_refs, "error": None}
    except Exception as e:
        logger.error(f"[{node_name}] Error formatting references: {e}")
        return {"error": f"Failed to format references: {str(e)}"}


async def format_complete_proposal(state: ExtendedProposalGraphState, config: RunnableConfig) -> Dict[str, Any]:
    """
    Node 7: Assembles all generated content into the final, formatted proposal document.
    """
    node_name = "format_complete_proposal"
    logger.info(f"--- Executing Node: {node_name} ---")
    state["node_history"].append(node_name)

    format_type = state.get("format_type", "ieee").upper()
    page_target = state.get("page_target", 7)

    # Consolidate all content pieces from the state
    content = {
        "Executive Summary": state.get("executive_summary", "Summary not generated."),
        "1. Introduction": {
            "1.1 Problem Statement": state.get("problem_statement", "Not generated."),
            "1.2 Background": state.get("problem_background", "Not generated."),
            "1.3 Significance": state.get("project_significance", "Not generated."),
            "1.4 Research Questions": "\n- ".join(state.get("research_questions", [])),
            "1.5 Project Objectives": "\n- ".join(state.get("project_objectives", [])),
        },
        "2. Literature Review": {
            "2.1 Summary": state.get("literature_summary", "Not generated."),
             **(state.get("literature_detailed_review", {})), # Unpack detailed sections
             "2.x Research Gaps": "\n- ".join(state.get("research_gaps", [])) # Adjust section number as needed
        },
        "3. Methodology": {
             "3.1 Overview": state.get("methodology_approach_overview", "Not generated."),
             **(state.get("methodology_sections", {})), # Unpack detailed sections
             "3.x Validation Strategy": state.get("validation_strategy", "Not generated."),
             "3.y Tools & Technologies": ", ".join(state.get("methodology_tools", [])),
        },
        "4. Expected Results and Discussion": {
            "4.1 Expected Results Summary": state.get("expected_results_summary", "Not generated."),
             **(state.get("expected_results_detailed", {})), # Unpack detailed sections
             **(state.get("discussion_sections", {})), # Unpack detailed sections
             "4.x Expected Impact": state.get("expected_impact", "Not generated."),
             "4.y Metrics": ", ".join(state.get("expected_metrics", []))
        },
        "5. Project Plan": {
            "5.1 Timeline & Milestones": state.get("timeline_and_milestones", "Not generated."),
            "5.2 Limitations & Future Work": state.get("limitations_and_future_work", "Not generated.")
        },
        "References": state.get("formatted_references", "References not generated.")
    }

    # Basic function to structure content for the prompt
    def format_content_for_prompt(data, indent=0):
        text = ""
        prefix = "  " * indent
        if isinstance(data, dict):
            for key, value in data.items():
                text += f"{prefix}**{key}**\n\n"
                text += format_content_for_prompt(value, indent + 1) + "\n"
        elif isinstance(data, list):
             text += f"{prefix}- " + f"\n{prefix}- ".join(data) + "\n\n" # Simple list format
        else:
            text += f"{prefix}{str(data)}\n\n"
        return text

    content_str = format_content_for_prompt(content)

    formatting_prompt_template = """You are an expert academic editor formatting a Final Year Project proposal according to {format_type} standards.
The target length is approximately {page_target} pages (this translates to roughly {word_target} words, assuming 300-400 words/page).
Assemble the following draft content into a single, coherent, and professionally formatted document.

Ensure:
- Correct {format_type} section headings and numbering (e.g., I. INTRODUCTION, II. LITERATURE REVIEW...).
- Formal academic tone and language.
- Logical flow between sections and paragraphs.
- Consistent formatting for lists, etc.
- Expand slightly or adjust wording for better flow and to meet the target length, but stay true to the core content.
- The final output should be the complete proposal text ONLY. Start with a Title Page placeholder if appropriate for the format.

**DRAFT CONTENT:**
{draft_content}

Begin the formatted proposal now:
"""
    # Estimate word target
    word_target = page_target * 350

    prompt = ChatPromptTemplate.from_template(formatting_prompt_template)
    chain = prompt | llm | StrOutputParser()

    try:
        final_proposal = await chain.ainvoke({
            "format_type": format_type,
            "page_target": page_target,
            "word_target": word_target,
            "draft_content": content_str
        }, config=config)
        logger.info(f"[{node_name}] Final proposal formatted successfully ({len(final_proposal.split())} words).")
        return {"formatted_proposal": final_proposal, "error": None}
    except Exception as e:
        logger.error(f"[{node_name}] Error formatting final proposal: {e}")
        return {"error": f"Failed to format final proposal: {str(e)}"}

# --- Graph Conditional Logic ---
def check_for_errors(state: ExtendedProposalGraphState) -> str:
    """Router function to check for errors and determine the next step."""
    if state.get("error"):
        logger.error(f"Graph execution failed with error: {state['error']}")
        return "error_handler" # Route to a dedicated error handler node or END

    last_node = state.get("node_history", [])[-1] if state.get("node_history") else None
    logger.debug(f"Routing check after node: {last_node}")

    # Define the standard flow
    node_sequence = [
        "initialize_project_context",
        "comprehensive_literature_review",
        "detailed_methodology_design",
        "generate_results_discussion_and_support",
        "generate_executive_summary",
        "format_references_node",
        "format_complete_proposal"
    ]

    if not last_node:
        return node_sequence[0] # Should not happen if entry point is set

    try:
        current_index = node_sequence.index(last_node)
        if current_index + 1 < len(node_sequence):
            next_node = node_sequence[current_index + 1]
            logger.debug(f"Routing from {last_node} to {next_node}")
            return next_node
        else:
            # Last node in sequence completed successfully
            logger.debug(f"Routing from {last_node} to END")
            return END
    except ValueError:
        # If last_node is not in the main sequence (e.g., error handler)
        logger.warning(f"Node {last_node} not in standard sequence. Ending.")
        return END

# Optional: Error Handler Node (can log, notify, or attempt recovery)
def handle_error(state: ExtendedProposalGraphState) -> Dict[str, Any]:
     node_name = "handle_error"
     logger.error(f"--- Executing Node: {node_name} ---")
     error_message = state.get("error", "Unknown error")
     logger.error(f"Proposal generation failed. Error: {error_message}")
     # You could add notification logic here
     # Returning the error message to potentially be included in the final state
     return {"formatted_proposal": f"PROPOSAL GENERATION FAILED: {error_message}"}

# --- Workflow Construction ---
workflow = StateGraph(ExtendedProposalGraphState)

# Add nodes
workflow.add_node("initialize_project_context", initialize_project_context)
workflow.add_node("comprehensive_literature_review", comprehensive_literature_review)
workflow.add_node("detailed_methodology_design", detailed_methodology_design)
workflow.add_node("generate_results_discussion_and_support", generate_results_discussion_and_support)
workflow.add_node("generate_executive_summary", generate_executive_summary)
workflow.add_node("format_references_node", format_references_node)
workflow.add_node("format_complete_proposal", format_complete_proposal)
workflow.add_node("error_handler", handle_error) # Add error handler node


# Define edges and entry point
workflow.set_entry_point("initialize_project_context")

# Use the single routing function after each node
workflow.add_conditional_edges("initialize_project_context", check_for_errors)
workflow.add_conditional_edges("comprehensive_literature_review", check_for_errors)
workflow.add_conditional_edges("detailed_methodology_design", check_for_errors)
workflow.add_conditional_edges("generate_results_discussion_and_support", check_for_errors)
workflow.add_conditional_edges("generate_executive_summary", check_for_errors)
workflow.add_conditional_edges("format_references_node", check_for_errors)
workflow.add_conditional_edges("format_complete_proposal", check_for_errors) # Routes to END if successful


# Add edge from error handler to END
workflow.add_edge("error_handler", END)


# Compile the graph
extended_proposal_app = workflow.compile()


# --- API Endpoint Implementation (Async) ---
proposal_jobs: Dict[str, Dict[str, Any]] = {} # In-memory job store

# async def run_extended_graph(job_id: str, initial_state: ExtendedProposalGraphState):
#     """Runs the extended graph asynchronously."""
#     proposal_jobs[job_id]["status"] = "running"
#     logger.info(f"Starting background graph execution for job {job_id}")
#     try:
#         config = {"configurable": {"thread_id": f"extended-proposal-{job_id}"}}
#         # Use ainvoke for the entire graph run
#         final_state = await extended_proposal_app.ainvoke(initial_state, config=config)

#         # Check final state for errors explicitly, even if routed through handler
#         final_error = final_state.get("error")
#         final_proposal = final_state.get("formatted_proposal")

#         if final_error or (not final_proposal) or ("FAILED" in str(final_proposal)):
#             status = "failed"
#             message = f"Proposal generation failed: {final_error or 'Unknown error during finalization.'}"
#             result = {"error": final_error, "node_history": final_state.get("node_history", [])}
#             logger.error(f"Job {job_id} failed: {message}")
#         else:
#             status = "completed"
#             message = f"Proposal generated successfully (~{len(final_proposal.split())} words)."
#             # Selectively return final pieces
#             result = {
#                 "formatted_proposal": final_proposal,
#                 "executive_summary": final_state.get("executive_summary"),
#                 "node_history": final_state.get("node_history", [])
#             }
#             logger.info(f"Job {job_id} completed successfully.")

#         proposal_jobs[job_id].update({
#             "status": status,
#             "message": message,
#             "result": result
#         })

#     except Exception as e:
#         logger.error(f"Unhandled exception during background graph execution for job {job_id}: {e}", exc_info=True)
#         proposal_jobs[job_id].update({
#             "status": "failed",
#             "message": f"An unexpected server error occurred: {str(e)}",
#             "result": {"error": str(e)}
#         })


# @router.post("/generate_extended_proposal", status_code=202, response_model=ExtendedProposalResponse)
# async def create_extended_proposal_job(
#     background_tasks: BackgroundTasks,
#     body: ExtendedProjectInput
# ):
#     """
#     Asynchronously starts the extended project proposal generation.
#     """
#     job_id = str(uuid.uuid4())
#     logger.info(f"Received extended proposal request for field='{body.field}', domain='{body.domain}'. Job ID: {job_id}")

#     # Initialize state dictionary from input
#     initial_state = ExtendedProposalGraphState(
#         field=body.field,
#         domain=body.domain,
#         idea=body.idea,
#         academic_level=body.academic_level,
#         page_target=body.page_target,
#         format_type=body.format_type,
#         # Initialize all other state keys to None or empty defaults
#         problem_statement=None, problem_background=None, project_significance=None,
#         key_challenges=None, search_queries=None, search_results=None,
#         literature_summary=None, literature_detailed_review=None, research_gaps=None,
#         research_questions=None, project_objectives=None,
#         methodology_approach_overview=None, methodology_sections=None, experimental_design=None,
#         methodology_tools=None, validation_strategy=None,
#         expected_results_summary=None, expected_results_detailed=None, discussion_sections=None,
#         expected_metrics=None, expected_impact=None, limitations_and_future_work=None,
#         timeline_and_milestones=None, references_data=None, formatted_references=None,
#         executive_summary=None, formatted_proposal=None, error=None, node_history=[]
#     )

#     # Store initial job info
#     proposal_jobs[job_id] = {
#         "proposal_id": job_id,
#         "status": "pending",
#         "message": "Extended proposal generation initiated.",
#         "result": None
#     }

#     # Run the graph in the background
#     background_tasks.add_task(run_extended_graph, job_id, initial_state)

#     return ExtendedProposalResponse(
#         proposal_id=job_id,
#         status="pending",
#         message="Extended proposal generation started. Check status using the job ID."
#     )

# @router.get("/generate_extended_proposal/status/{job_id}", response_model=ExtendedProposalResponse)
# async def get_extended_proposal_status(job_id: str):
#     """
#     Checks the status and result of an extended proposal generation job.
#     """
#     job_info = proposal_jobs.get(job_id)
#     if not job_info:
#         raise HTTPException(status_code=404, detail="Job ID not found.")

#     # Ensure the response matches the Pydantic model structure
#     return ExtendedProposalResponse(**job_info)







class ProjectIdeaInput(BaseModel):
    area_of_interest: List[str] = Field(..., description="The area of interest for the project.")
    academic_level: str = Field(..., description="The academic level of the project.")
    skills: List[str] = Field(..., description="The skills of the student.")
    preferred_project_type: List[str] = Field(..., description="Students indicate whether they prefer a research-oriented project (e.g., studying a new algorithm), a development-oriented project (e.g., building an application), or a hybrid of both")
    real_world_problem_or_inspiration: Optional[str] = Field(..., description="Students can optionally describe a specific problem they've encountered or an idea they'd like to explore, even if it's vague (e.g., 'I want to improve online learning tools')")
    project_duration: str = Field(..., description="The duration of the project.")
    team_size: str = Field(..., description="The size of the team.")


class RoadMapInput(BaseModel):
    project_id: str = Field(..., description="The ID of the project to generate roadmap for")


class RoadMapResponse(BaseModel):
    roadmap_id: str
    status: str
    message: str
    result: str = None


from app.api.routes.problem_identification import problem_conversation
from app.api.routes.history_schema import History
@router.post("/generate_project_idea")
async def generate_project_idea(body: ProjectIdeaInput, history: History):
    """
    Generates a project idea based on the user's input.
    """
    if not history.history:
        history.history = [{
            "role": "human",
            "content": f"I am a student of {body.academic_level} level and my area of interest is {body.area_of_interest} and my skills are {body.skills} and my preferred project type is {body.preferred_project_type} and my real world problem or inspiration is {body.real_world_problem_or_inspiration} and my project duration is {body.project_duration} and my team size is {body.team_size} want to find a project idea"
        }]
    print(history)
    problem_statement = problem_conversation(history.history)
    return {"history": problem_statement}


from docx import Document
from fastapi.responses import FileResponse
import os

@router.post("/generate_roadmap")
async def generate_roadmap(body: ProjectIdeaInput, project_id: RoadMapInput):
    """
    Generates a comprehensive roadmap for a project based on its problem statement.
    Saves the roadmap as a Word document with the project_id as filename.
    """
    # Load the problem statement
    file_path = "./users_ids.json"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="No problem statements found")
    
    with open(file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    
    # Find the problem statement with matching ID
    problem_statement = None
    for statement in data.get('problem_statements', []):
        if project_id.project_id in statement:
            problem_statement = statement[project_id.project_id]
            break
    
    if not problem_statement:
        raise HTTPException(status_code=404, detail="Project ID not found")
    
    # Generate roadmap using LLM
    roadmap_prompt = f"""
    Based on the following problem statement, create a comprehensive project roadmap:
    
    ## Student Details:
    Area of Interest: {body.area_of_interest}
    Academic Level: {body.academic_level}
    Skills: {body.skills}
    Preferred Project Type: {body.preferred_project_type}
    Real World Problem or Inspiration: {body.real_world_problem_or_inspiration}
    project Duration: {body.project_duration}
    Team Size: {body.team_size}
    
    ## Problem Statement:
    {problem_statement}
    
    The roadmap should include:
    1. Project phases with timelines
    2. Key milestones and deliverables
    3. Required resources
    4. Potential risks and mitigation strategies
    5. Success metrics
    """
    
    if not llm:
        raise HTTPException(status_code=500, detail="LLM not initialized")
    
    try:
        response = await llm.ainvoke(roadmap_prompt)
        roadmap_content = response.content
        
        # Save as Word document
        doc = Document()
        doc.add_heading(f'Project Roadmap - {project_id.project_id}', level=1)
        
        # Add student details section
        doc.add_heading('Student Details', level=2)
        doc.add_paragraph(f"Area of Interest: {', '.join(body.area_of_interest)}")
        doc.add_paragraph(f"Academic Level: {body.academic_level}")
        doc.add_paragraph(f"Skills: {', '.join(body.skills)}")
        doc.add_paragraph(f"Project Type: {', '.join(body.preferred_project_type)}")
        doc.add_paragraph(f"Duration: {body.project_duration}")
        doc.add_paragraph(f"Team Size: {body.team_size}")
        
        # Add roadmap content
        doc.add_heading('Project Roadmap', level=2)
        doc.add_paragraph(roadmap_content)
        
        # Create roadmaps directory if it doesn't exist
        os.makedirs('roadmaps', exist_ok=True)
        doc_path = f'roadmaps/{project_id.project_id}.docx'
        doc.save(doc_path)
        
        return RoadMapResponse(
            roadmap_id=project_id.project_id,
            status="completed",
            message="Roadmap generated and saved successfully",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate roadmap: {str(e)}")

@router.get("/download_roadmap/{project_id}")
async def download_roadmap(project_id: str):
    """
    Downloads the generated roadmap Word document.
    """
    doc_path = f'roadmaps/{project_id}.docx'
    if not os.path.exists(doc_path):
        raise HTTPException(status_code=404, detail="Roadmap not found")
    
    return FileResponse(
        doc_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=f'project_roadmap_{project_id}.docx'
    )

